#!/usr/bin/env python3
"""Build Chinese CNC605+ user manual PDF/DOCX from English source."""

from __future__ import annotations

import re
import subprocess
from pathlib import Path

import fitz
from docx import Document

ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "docs" / "CNC605+"
SRC_PDF = OUT_DIR / "CNC605+_User_Manual_EN.pdf"
OUT_PDF = OUT_DIR / "CNC605+_用户使用手册_CN.pdf"
OUT_DOCX = OUT_DIR / "CNC605+_用户使用手册_CN.docx"
FONT_PATH = "/usr/share/fonts/truetype/wqy/wqy-microhei.ttc"

# Page header/footer labels kept bilingual where appropriate
TITLE_EN = "CNC605+ User Manual"
TITLE_CN = "CNC605+ 用户使用手册"

# Ordered longest-first replacements for DOCX fuzzy matching
TRANSLATIONS: list[tuple[str, str]] = [
    (
        "All rights reserved! Any company or individual person shall not copy or back up this user manual in any format (electronic, mechanical, photocopying, recording or other formats) without written permission from LAUNCH TECH CO., LTD. (hereinafter referred to as \"Launch\"). The manual is for the use of the products manufactured by Launch, which shall not assume any responsibility for the consequences arising from the use of it to guide the operations of other equipment.",
        "版权所有！未经深圳市元征科技股份有限公司（以下简称“元征”）书面许可，任何公司或个人不得以任何形式（电子、机械、影印、录音或其他形式）复制或备份本用户手册。本手册仅适用于元征制造的产品，对于使用本手册指导其他设备操作所产生的后果，元征不承担任何责任。",
    ),
    (
        "Launch and its branches will not bear any liability for the fees and expenses incurred by equipment damage or loss due to accidents caused by users or third parties, misuses and abuses, unauthorized modifications and repairs, or operations and services not following Launch's instructions.",
        "对于因用户或第三方造成的事故、误用和滥用、未经授权的改装和维修，或未按元征说明进行操作和服务而导致的设备损坏或损失所产生的费用和开支，元征及其分支机构不承担任何责任。",
    ),
    (
        "Launch assumes no responsibility for device damages or problems resulted from the usage of other parts or consumables, rather than original products of Launch or products approved by the company.",
        "对于因使用非元征原装产品或未经公司认可的产品之外的零件或耗材而导致的设备损坏或故障，元征不承担任何责任。",
    ),
    (
        "Official statement: The mentioning of the names of other products in this manual is to illustrate how to use the device, with the ownership of the registered trademarks belonging to the owners.",
        "特别声明：本手册中提及其他产品名称，仅为说明本设备的使用方法，相关注册商标所有权归各权利人所有。",
    ),
    (
        "The device is intended for the use of professional technicians or maintenance and repair personnel.",
        "本设备仅供专业技师或维修保养人员使用。",
    ),
    ("Registered Trademark", "注册商标"),
    (
        "Launch has registered its trademark in China and several other countries, and the logo is . In those countries where trademarks, service marks, domain names, icons, company names of Launch have not been registered yet, Launch declaims the right for its unregistered trademarks, service marks, domain names, icons, and company names. Trademarks of other products and company names mentioned in this manual are still owned by the original registered companies. Without written agreement from the owner, no person is allowed to use the trademarks, service marks, domain names, icons and company names of Launch or of other mentioned companies. You can visit https://www.cnlaunch.com, or write to Customer Service Center of LAUNCH TECH CO., LTD. at No.4012, Launch Industrial Park, North Wuhe Rd, Bantian Street, Longgang District, Shenzhen, China, to get contact with Launch for the written agreement on the usage of the user manual.",
        "元征已在中国及多个国家注册其商标，标识为 。在元征商标、服务标志、域名、图标、公司名称尚未注册的国家/地区，元征保留对其未注册商标、服务标志、域名、图标及公司名称的权利。本手册中提及的其他产品及公司名称之商标，仍归原注册公司所有。未经权利人书面同意，任何人不得使用元征或其他相关公司的商标、服务标志、域名、图标及公司名称。如需就本用户手册的使用与元征取得联系并获取书面授权，请访问 https://www.cnlaunch.com，或致函：深圳市元征科技股份有限公司客户服务中心，地址：中国深圳市龙岗区坂田街道五和大道北元征工业园4012号。",
    ),
    ("Disclaimer of Warranties and Limitation of Liabilities", "免责声明与责任限制"),
    (
        "All information, illustrations, and specifications in this manual are based on the latest information available at the time of publication.",
        "本手册中的所有信息、图示及规格均基于出版时可获得的最新资料。",
    ),
    (
        "The right is reserved to make changes at any time without notice. We shall not be liable for any direct, special, incidental, indirect damages or any economic consequential damages (including the loss of profits) due to the use of the document.",
        "我们保留随时更改的权利，恕不另行通知。对于因使用本文件而造成的任何直接、特殊、附带、间接损害或任何经济后果性损害（包括利润损失），我们不承担责任。",
    ),
    (
        "Special note: The manual introduces in detail the structure, functions, operation method and related accessories of injector cleaner & tester, and briefly explains the precautions in the process of using, maintenance and handling of abnormalities. Launch reserves the right to change the design and specifications of the product. The actual configuration shall be subject to the packing list.",
        "特别说明：本手册详细介绍喷油嘴清洗检测仪的结构、功能、操作方法及相关附件，并简要说明使用、维护及异常处理过程中的注意事项。元征保留更改产品设计和规格的权利，实际配置以装箱单为准。",
    ),
    ("ATTENTION！", "注意！"),
    ("Important Safety Instructions:", "重要安全须知："),
    ("Please read the manual carefully before using the instrument for proper operation!", "使用仪器前请仔细阅读本手册，以确保正确操作！"),
    ("Be careful when touching the device or hot parts of the engine.", "触摸设备或发动机高温部件时请小心。"),
    (
        "If the power cord is broken, please do not turn the device on and use it. If the device is dropped or damaged, please use it after being inspected by a professional.",
        "如电源线破损，请勿开机使用。如设备跌落或损坏，须经专业人员检查后方可使用。",
    ),
    (
        "Please do not hang the power cord on the edge of the table, chair or counter. Do not touch hot parts or rotating fan blades.",
        "请勿将电源线挂在桌边、椅边或柜台边缘。请勿触摸高温部件或旋转风扇叶片。",
    ),
    (
        "If it is necessary to extend the power cable, the level of the power cord should be higher than or equal to that of the original power line. Overheat may occur if a power cable of inferior level is used.",
        "如需延长电源线，延长线的规格等级应不低于原电源线。使用低规格电源线可能导致过热。",
    ),
    (
        "Do not connect the plug of power cord when not using the device. Do not remove the plug by pulling the power line. It should be unplugged by hand.",
        "不使用设备时请勿连接电源插头。拔插头时请勿拉扯电源线，应手动拔出。",
    ),
    ("The device should be cooled completely before storage and the line should be wound up.", "存放前设备应完全冷却，并将电源线收好。"),
    (
        "The cleaning agent for the device is a flammable and weakly volatile liquid. Smoking and lighting fires are strictly prohibited during the cleaning process.",
        "设备所用清洗液为易燃、弱挥发性液体。清洗过程中严禁吸烟和明火。",
    ),
    (
        "The instrument should be placed in a room that is not exposed to direct sunlight and is well ventilated, and signs \"Smoking and lighting fires strictly forbidden\" and \"Danger warning of inflammables\" should be posted.",
        "仪器应放置在避免阳光直射、通风良好的房间内，并张贴「严禁烟火」和「易燃危险品警告」标识。",
    ),
    (
        "The operator's hair, clothes, fingers and other parts of the body should be kept away from the operating parts of the equipment.",
        "操作人员的头发、衣物、手指及身体其他部位应远离设备运转部件。",
    ),
    (
        "To prevent electric shocks, do you touch the operating equipment in wet areas or operate it in the rain.",
        "为防止触电，请勿在潮湿环境或雨中触摸或操作设备。",
    ),
    (
        "Please operate the device as described in the manual. Use accessories recommended by the manufacturer.",
        "请按本手册说明操作设备，并使用制造商推荐的附件。",
    ),
    (
        "It is strictly forbidden to open the ultrasonic system under the circumstances that ultrasonic cleaning agent has not been added into the ultrasonic cleaning pool. Otherwise, the ultrasonic equipment may be damaged easily.",
        "严禁在未向超声波清洗槽加入超声波清洗剂的情况下启动超声波系统，否则超声波设备极易损坏。",
    ),
    ("The housing of the device must be reliable and grounded.", "设备外壳必须可靠接地。"),
    (
        "Automobile exhaust contains a variety of toxic and harmful gases (such as carbon monoxide, hydrocarbon, nitrogen oxide and etc.). During the test, the exhaust should be directed outdoors and the room should be well ventilated.",
        "汽车尾气含有多种有毒有害气体（如一氧化碳、碳氢化合物、氮氧化物等）。测试时应将尾气排至室外，并保持室内通风良好。",
    ),
    (
        "The temperature of the exhaust pipe and water tank of the automobile's engine is high. Do not touch them to prevent burns.",
        "汽车发动机排气管和水箱温度很高，请勿触摸，以免烫伤。",
    ),
    (
        "Please pull up the handbrake of the vehicle to be cleaned, shift the transmission to neutral position and block the front wheel before free disassembly cleaning.",
        "进行免拆清洗前，请拉紧待清洗车辆的手刹，将变速箱挂入空挡，并楔住前轮。",
    ),
    ("Wear safety glasses when operating. Daily wear glasses are not safety glasses.", "操作时请佩戴防护眼镜。普通眼镜不是安全眼镜。"),
    (
        "When disconnecting a pressurized fuel pipe union, please cover the union with a towel to avoid getting hurt by fuel gushing out and causing fires.",
        "拆卸加压燃油管接头时，请用毛巾盖住接头，防止燃油喷出伤人或引发火灾。",
    ),
    (
        "Test solution is used by the main unit of the device uses, and ultrasonic cleaning agent is used for ultrasonic cleaning.",
        "主机使用检测液，超声波清洗使用超声波清洗剂。",
    ),
    (": Indicates where attention should be paid when operating the device.", "：表示操作设备时应注意之处。"),
    (": Indicates the possibility of product damage and personnel injury during operation.", "：表示操作过程中可能导致产品损坏或人员受伤。"),
    ("Content", "目录"),
    (
        "I. Introduction to Injector Cleaner & Tester.....................................................................................................1 1.1 Brief Introduction....................................................................................................................................... 1 1.2 Functions and Features.............................................................................................................................1 1.3 Working Environment and Specifications.................................................................................................. 2 II. Structure of Injector Cleaner & Tester.........................................................................................................2 2.1 Structure....................................................................................................................................................2 2.2 Control Panel.............................................................................................................................................3 III. Installation and Connection........................................................................................................................ 3 3.1 Installation................................................................................................................................................. 3 3.2 Connection................................................................................................................................................ 3 IV. Operation Procedures for Cleaning and Testing of Injectors...................................................................3 4.1 Preparations.............................................................................................................................................. 3 4.2 Cleaning and Testing Sequences.............................................................................................................. 4 4.3 Cleanup after Operation............................................................................................................................ 4 V. Operation Processes.................................................................................................................................... 5 5.1 Operation Mode Selection.........................................................................................................................5 5.2 System Setting.......................................................................................................................................... 5 5.3 Resistance Test......................................................................................................................................... 7 5.4 Ultrasonic Cleaning................................................................................................................................... 7 5.5 Leakage Test............................................................................................................................................. 8 5.6 Uniformity/Spray........................................................................................................................................9 5.7 Injection Quantity.....................................................................................................................................11 5.8 Auto Mode................................................................................................................................................11 5.9 On-Vehicle Clean.....................................................................................................................................13 VI. Service and Maintenance..........................................................................................................................16 6.1 Handling, Storage and Installation Environment......................................................................................16 6.2 Common Quick-Wear Parts and Consumables.......................................................................................16 6.3 Precautions and Solutions to Common Problems................................................................................... 17 Appendix I: Fuel System Pressure Gauge of Typical Vehicles....................................................................18",
        "一、喷油嘴清洗检测仪简介.................................................................................................................................1\n1.1 简要介绍...................................................................................................................................................1\n1.2 功能与特点...............................................................................................................................................1\n1.3 工作环境与技术规格................................................................................................................................2\n二、喷油嘴清洗检测仪结构...........................................................................................................................2\n2.1 结构..........................................................................................................................................................2\n2.2 控制面板...................................................................................................................................................3\n三、安装与连接..............................................................................................................................................3\n3.1 安装..........................................................................................................................................................3\n3.2 连接.........................................................................................................................................................3\n四、喷油嘴清洗与检测操作规程...............................................................................................................3\n4.1 准备工作..................................................................................................................................................3\n4.2 清洗与检测顺序......................................................................................................................................4\n4.3 操作后清理................................................................................................................................................4\n五、操作流程..................................................................................................................................................5\n5.1 工作模式选择...........................................................................................................................................5\n5.2 系统设置...................................................................................................................................................5\n5.3 电阻检测...................................................................................................................................................7\n5.4 超声波清洗.............................................................................................................................................7\n5.5 泄漏检测.................................................................................................................................................8\n5.6 均匀性/雾化检测......................................................................................................................................9\n5.7 喷油量检测.............................................................................................................................................11\n5.8 自动模式.................................................................................................................................................11\n5.9 免拆清洗.................................................................................................................................................13\n六、保养与维护..........................................................................................................................................16\n6.1 搬运、存放与安装环境........................................................................................................................16\n6.2 常用易损件与耗材...............................................................................................................................16\n6.3 注意事项与常见问题处理...................................................................................................................17\n附录一：典型车型燃油系统压力参考表...................................................................................................18",
    ),
    ("Introduction to Injector Cleaner & Tester", "喷油嘴清洗检测仪简介"),
    ("1.1 Brief Introduction", "1.1 简要介绍"),
    ("1.2 Functions and Features", "1.2 功能与特点"),
    ("1.3 Working Environment and Specifications", "1.3 工作环境与技术规格"),
    ("II. Structure of Injector Cleaner & Tester", "二、喷油嘴清洗检测仪结构"),
    ("2.1 Structure", "2.1 结构"),
    ("2.2 Control Panel", "2.2 控制面板"),
    ("III. Installation and Connection", "三、安装与连接"),
    ("3.1 Installation", "3.1 安装"),
    ("3.2 Connection", "3.2 连接"),
    ("IV. Operation Procedures for Cleaning and Testing of Injectors", "四、喷油嘴清洗与检测操作规程"),
    ("4.1 Preparations", "4.1 准备工作"),
    ("4.2 Cleaning and Testing Sequences", "4.2 清洗与检测顺序"),
    ("4.3 Cleanup after Operation", "4.3 操作后清理"),
    ("V. Operation Processes", "五、操作流程"),
    ("5.1 Operation Mode Selection", "5.1 工作模式选择"),
    ("5.2 System Setting", "5.2 系统设置"),
    ("5.3 Resistance Test", "5.3 电阻检测"),
    ("5.4 Ultrasonic Cleaning", "5.4 超声波清洗"),
    ("5.5 Leakage Test", "5.5 泄漏检测"),
    ("5.6 Uniformity/Spray", "5.6 均匀性/雾化检测"),
    ("5.7 Injection Quantity", "5.7 喷油量检测"),
    ("5.8 Auto Mode", "5.8 自动模式"),
    ("5.9 On-Vehicle Clean", "5.9 免拆清洗"),
    ("VI. Service and Maintenance", "六、保养与维护"),
    ("6.1 Handling, Storage and Installation Environment", "6.1 搬运、存放与安装环境"),
    ("6.2 Common Quick-Wear Parts and Consumables", "6.2 常用易损件与耗材"),
    ("6.3 Precautions and Solutions to Common Problems", "6.3 注意事项与常见问题处理"),
    ("Appendix I: Fuel System Pressure Gauge of Typical Vehicles", "附录一：典型车型燃油系统压力参考表"),
    (
        "Welcome to use Injector Cleaner & Tester manufactured by LAUNCH TECH CO., LTD. It is a mechatronic product that combines ultrasonic cleaning technology with microprocessor oil pressure control cleaning and testing technology. It is capable to simulate various working conditions of engine and perform cleaning and testing for injector.",
        "欢迎使用深圳市元征科技股份有限公司制造的喷油嘴清洗检测仪。本产品为机电一体化设备，集超声波清洗技术与微处理器油压控制清洗检测技术于一体，可模拟发动机多种工况，对喷油嘴进行清洗与检测。",
    ),
    (
        "The User Manual is applicable to the following product:",
        "本用户手册适用于以下产品：",
    ),
    (
        "Injector Cleaner & Tester: a desk 6-cylinder injector cleaner and tester.",
        "喷油嘴清洗检测仪：台式六缸喷油嘴清洗检测仪。",
    ),
    ("Main Functions", "主要功能"),
    (
        "Working mode selection: Select EFI, GDI and PIEZ0 operating modes according to the injector types.",
        "工作模式选择：根据喷油嘴类型选择 EFI（12V）、GDI（65V）或 PIEZO（140V）工作模式。",
    ),
    ("Set: To configure the system settings of the device.", "设置：配置设备系统参数。"),
    (
        "Resistance test: Through the pulse signal line, the internal resistance of up to six injectors can be tested simultaneously to determine whether the injector circuit condition is normal.",
        "电阻检测：通过脉冲信号线，可同时检测最多六个喷油嘴的内部电阻，判断喷油嘴电路是否正常。",
    ),
    (
        "Ultrasonic clean: To perform ultrasonic cleaning on multiple injectors at the same time, removing carbon deposits on injectors completely.",
        "超声波清洗：可同时对多个喷油嘴进行超声波清洗，彻底清除喷油嘴积碳。",
    ),
    (
        "Leakage test: To test the leakage and dribbling conditions of injectors under system pressure.",
        "泄漏检测：在系统压力下检测喷油嘴的泄漏及滴漏情况。",
    ),
    (
        "Uniformity/Spray: To detect the uniformity of the fuel injection quantity of each injector and monitor the spray status of injectors thoroughly and carefully by using the backlight, and to reverse flush injectors.",
        "均匀性/雾化检测：检测各喷油嘴喷油量的均匀性，并利用背景灯全面仔细观察喷油雾化状态，同时可对喷油嘴进行反向冲洗。",
    ),
    (
        "Injection quantity: To detect the amount of fuel injected normally by injectors in 20 seconds.",
        "喷油量检测：检测喷油嘴在20秒常喷条件下的正常喷油量。",
    ),
    (
        "Auto Mode: under specific working conditions and parameters, the precise simulation of test of injectors under various working conditions",
        "自动模式：在特定工况参数下，真实模拟喷油嘴在各种工况下的检测。",
    ),
    (
        "On-vehicle clean: Coming with a variety of disassembly-free cleaning connectors, the device can be used to perform On-vehicle clean and maintenance for various vehicle models.",
        "免拆清洗：配备多种免拆清洗接头，可对多种车型进行免拆清洗维护。",
    ),
    ("Note: This function is only available for EFI working mode.", "注意：此功能仅在 EFI 工作模式下可用。"),
    ("Main Features", "主要特点"),
    (
        "Adopting ultrasonic cleaning technology, Injector Cleaner & Tester presents a strong cleaning ability;",
        "采用超声波清洗技术，清洗能力强；",
    ),
    (
        "The product also adopts fuel pressure adjustment and control technology by microcomputer, which can ensure stable fuel pressure and wide adjustable range. It is applicable to vehicles equipped with a variety of gasoline injection systems. Meanwhile, the automation of injectors' cleaning and testing processes can be realized.",
        "产品采用微电脑燃油压力调节与控制技术，可确保燃油压力稳定、调节范围宽，适用于配备多种汽油喷射系统的车辆，同时可实现喷油嘴清洗与检测过程的自动化。",
    ),
    (
        "Thanks to the adoption of microcomputer automatic control and digital display technologies, the cleaning and testing processes can be controlled automatically and the parameters of the main status can be monitored in real time.",
        "采用微电脑自动控制及数字显示技术，可自动控制清洗检测过程，并实时监控主要状态参数。",
    ),
    ("Specifications:", "技术规格："),
    ("Working Environment:", "工作环境："),
    ("Power: AC110V, 50Hz/60Hz", "电源：AC110V，50Hz/60Hz"),
    ("AC220V, 50Hz/60Hz", "AC220V，50Hz/60Hz"),
    ("Mechanical Power: 500W", "整机功率：500W"),
    ("Fuel Tank Capacity: 2500ml", "油箱容量：2500ml"),
    ("Ultrasonic Cleaning Machine Power: 100W", "超声波清洗机功率：100W"),
    ("Working Pressure: 0~9bar", "工作压力：0~9bar"),
    ("Simulation Test Speed Range: 100~9900rpm (step length: 10rpm)", "模拟转速范围：100~9900rpm（步进：10rpm）"),
    ("Timing Range: 1~9999s", "计时范围：1~9999s"),
    ("Pulse Width Range: 0.1~25ms (step length: 0.1ms)", "脉宽范围：0.1~25ms（步进：0.1ms）"),
    ("Resistance Test Range: 0~250Ω", "电阻检测范围：0~250Ω"),
    ("Environment Temperature: 0℃~45℃", "环境温度：0℃~45℃"),
    ("Relative Humidity: ＜85%", "相对湿度：＜85%"),
    ("Strength of External Magnetic Field: ＜400A/m", "外磁场强度：＜400A/m"),
    ("Open fires is strictly prohibited within 2m.", "2米范围内严禁明火。"),
    (
        "The schematic diagram of Injector Cleaner & Tester is shown in Figure 2.1:",
        "喷油嘴清洗检测仪结构示意图见图2.1：",
    ),
    (
        "Attention: There may be slight difference between the illustrations in this manual and the actual product. The actual product prevails.",
        "注意：本手册图示与实际产品可能略有差异，以实际产品为准。",
    ),
    ("Figure 2.1 Schematic diagram", "图2.1 结构示意图"),
    (
        "1-Oil supply pipe; 2- Ultrasonic cleaner; 3- Oil tank liquid level pipe+Oil tank drain pipe; 4- Pulse signal line; 5-Uniformity/Spray Observation Glass; 6- Control panel (operation buttons and knobs); 7- Pressure Gauge; 8- Fuel distributor assembly; 9-Pulse signal line fixed position; 10- Power switch; 11- Power Socket ; 12- Filter; 13- Oil return port; 14- Fuel pump; 15- Fuel tank",
        "1-供油管；2-超声波清洗机；3-油箱液位管+油箱放油管；4-脉冲信号线；5-均匀性/雾化观察玻璃管；6-控制面板（操作按钮与旋钮）；7-压力表；8-燃油分配器总成；9-脉冲信号线固定位置；10-电源开关；11-电源插座；12-过滤器；13-回油口；14-燃油泵；15-油箱",
    ),
    ("The control panel is shown in Figure 2.2:", "控制面板见图2.2："),
    ("Figure 2.2 Diagram of control panel", "图2.2 控制面板示意图"),
    ("The installation steps are as follows:", "安装步骤如下："),
    ("1）Move the packaged machine to a flat surface.", "1）将包装好的机器移至平整地面。"),
    (
        "2）Check if the packaging, machine, accessory box, user manual, power cord and etc. are complete.",
        "2）检查包装、主机、附件箱、用户手册、电源线等是否齐全。",
    ),
    (
        "Take the power cord out of the packaging box and plug it into the power input port on the right side of the machine.",
        "从包装箱中取出电源线，插入机器右侧电源输入口。",
    ),
    (
        "1）Remove the injector from the vehicle and check if the rubber seal of the injector is damaged. If it is damaged, replace it with a seal of the same type before cleaning and testing to avoid leak. Put the injector into gasoline or detergent, carefully remove the greasy dirt outside and then wipe it with a soft cloth.",
        "1）从车辆上拆下喷油嘴，检查喷油嘴橡胶密封圈是否损坏。如已损坏，清洗检测前应更换同型号密封圈，以防泄漏。将喷油嘴放入汽油或清洗剂中，小心清除外部油污，再用软布擦拭干净。",
    ),
    (
        "2）Observe the test solution level to ensure that there is enough test solution in the tank. Pour the test solution into the tank from the filling port on the upper left corner of the device and observe the level through the level display tube. In most cases, fill 2000ml of test solution into the fuel filling port. Pay attention to the filled test solution shall not exceed the warning line.",
        "2）观察检测液液位，确保油箱内有足够检测液。从设备左上角加注口倒入检测液，通过液位显示管观察液位。一般情况下，向燃油加注口加注2000ml检测液。注意：加注量不得超过警告线。",
    ),
    ("3）Turn on the power switch on the right of the cabinet.", "3）打开机柜右侧电源开关。"),
    (
        "4）Add an appropriate amount of cleaning agent into the ultrasonic cleaning tank to immerse the needle valve of the injector.",
        "4）向超声波清洗槽中加入适量清洗剂，使喷油嘴针阀浸没在清洗液中。",
    ),
    (
        "Pour ultrasonic detergent into the ultrasonic cleaning basin so that the needle valve of the injector is covered by the detergent.",
        "向超声波清洗槽倒入超声波清洗剂，使喷油嘴针阀被清洗剂完全覆盖。",
    ),
    (
        "5）Put the injector into ultrasonic cleaning tank (the end connected to the wire is facing upwards, and the pointed foot is facing downwards).",
        "5）将喷油嘴放入超声波清洗槽（接线端朝上，针阀端朝下）。",
    ),
    (
        "Attention: The main unit of the device uses test solution for uniformity/spray test, leakage test, injection quantity test and auto. mode test. The ultrasonic cleaner uses the cleaning solution. Test solution and cleaning solution are not included in the standard configuration and can be purchased separately.",
        "注意：主机使用检测液进行均匀性/雾化检测、泄漏检测、喷油量检测及自动模式检测；超声波清洗机使用清洗液。检测液和清洗液均不含在标准配置中，可另行购买。",
    ),
    ("Figure 4.1", "图4.1"),
    (
        "1- Schematic diagram of the testing solution; 2- Schematic diagram of cleaning solution.",
        "1-检测液示意图；2-清洗液示意图。",
    ),
    (
        "It is recommended to carry out the complete cleaning and testing procedures in the following order.",
        "建议按以下顺序完成全套清洗与检测流程。",
    ),
    ("Resistance Test", "电阻检测"),
    ("Ultrasonic clean", "超声波清洗"),
    ("Leakage test", "泄漏检测"),
    ("Uniformity/Spray test", "均匀性/雾化检测"),
    ("Injecting quantity test", "喷油量检测"),
    ("Auto Mode test", "自动模式检测"),
    (
        "According to different test items, select the corresponding parameters and set them. See \"V. Operation Processes\" for details.",
        "根据不同检测项目选择并设置相应参数，详见“五、操作流程”。",
    ),
    (
        "After the end of cleaning and testing, a cleanup should be done, which includes:",
        "清洗与检测结束后应进行清理，包括：",
    ),
    (
        "Press the fuel drain button on the control panel to drain test solution to a fuel container.",
        "按控制面板上的排油按钮，将检测液排入燃油容器。",
    ),
    ("Switch off the power switch and unplug the power plug.", "关闭电源开关并拔掉电源插头。"),
    (
        "Pouring out all cleaning solution from the ultrasonic cleaning pool completely shown in Figure 4.2, and wipe the ultrasonic cleaning unit with a soft dry cloth.",
        "将超声波清洗槽内清洗液完全倒出（见图4.2），并用柔软干布擦拭超声波清洗单元。",
    ),
    ("Figure 4.2", "图4.2"),
    ("Wipe the table top of the machine with a soft dry cloth.", "用柔软干布擦拭机器台面。"),
    (
        "To avoid volatilization, drain all the test solution from the fuel tank. If it can be used again, store it in a safe place. If it is dirty and cannot be used any more, dispose of it according to relevant regulations.",
        "为避免挥发，应将油箱内检测液全部排出。如可继续使用，请妥善存放；如已污染不可再用，请按相关规定处置。",
    ),
    (
        "After turning on the device, the system will enter the operation mode selection interface as shown in the below figure. Please select the operating mode (EFI, GDI and PIEZO) according to the injector types .",
        "开机后，系统进入工作模式选择界面（见下图）。请根据喷油嘴类型选择工作模式（EFI、GDI 或 PIEZO）。",
    ),
    (
        "Different operation modes may have different workflows and parameters. Please confirm the injector type and select the correct operation mode, otherwise the injector may be damaged.",
        "不同工作模式的操作流程和参数可能不同。请确认喷油嘴类型并选择正确的工作模式，否则可能损坏喷油嘴。",
    ),
    ("Figure 5.1", "图5.1"),
    (
        "This function is used to change the system language, modify system parameters and view software version information.",
        "此功能用于更改系统语言、修改系统参数及查看软件版本信息。",
    ),
    ("Figure 5.2", "图5.2"),
    ("5.2.1 Language Selection", "5.2.1 语言选择"),
    (
        "In order to meet the needs of different countries and regions, multiple languages are available. Users can choose the appropriate language as needed.",
        "为满足不同国家和地区的需求，系统提供多种语言，用户可按需选择。",
    ),
    ("Figure 5.3", "图5.3"),
    ("5.2.2 Parameter", "5.2.2 参数"),
    (
        "Users can adjust screen brightness, system volume and restore factory settings as needed.",
        "用户可按需调节屏幕亮度、系统音量及恢复出厂设置。",
    ),
    (
        "Restoring factory settings will clear the parameters set by users, please operate with caution.",
        "恢复出厂设置将清除用户已设置的参数，请谨慎操作。",
    ),
    ("Figure 5.4", "图5.4"),
    ("5.2.3 Version", "5.2.3 版本"),
    ("Click [Version] button to check information of the current software version.", "点击【版本】按钮可查看当前软件版本信息。"),
    ("Figure 5.5", "图5.5"),
    (
        "This function is used to determine the quality of the fuel injectors by detecting their resistance.",
        "此功能通过检测喷油嘴电阻来判断喷油嘴质量。",
    ),
    ("Method and Steps:", "方法与步骤："),
    (
        "1）Connect the pulse signal line to the corresponding injector according to the color and serial number on the pulse signal line connectors and the fuel distributor.",
        "1）根据脉冲信号线接插件及燃油分配器上的颜色和序号，将脉冲信号线连接至对应喷油嘴。",
    ),
    (
        "2）Select [Resistance Test] on the main interface enter the below interface. Press [Start] button to start the test.",
        "2）在主界面选择【电阻检测】进入如下界面，按【开始】按钮启动检测。",
    ),
    (
        "3）The test results will be displayed below the corresponding serial number. If the fuel injector passes the test, its resistance value will be displayed below the corresponding serial number; If [SC] is displayed, it indicates a short circuit; and if [ERROR] is displayed, it indicates that there is a connection fault, poor contact, or excessive resistance issue with the fuel injector.",
        "3）检测结果将显示在对应序号下方。如喷油嘴检测合格，对应序号下方显示电阻值；如显示【SC】，表示短路；如显示【ERROR】，表示喷油嘴存在连接故障、接触不良或电阻过大。",
    ),
    ("Warning!", "警告！"),
    (
        "You MUST stop running the system before plugging or unplugging the pulse signal line.",
        "插拔脉冲信号线前，必须先停止系统运行。",
    ),
    (
        "Ultrasonic cleaning is an advanced cleaning method that uses the penetration and cavitation shock waves generated by the propagation of the ultrasonic waves in the medium to run a power cleaning on objects with complex shapes, cavities and pores, in order to remove stubborn carbon deposits on the injector thoroughly.",
        "超声波清洗是一种先进清洗方法，利用超声波在介质中传播产生的穿透力和空化冲击波，对形状复杂、有空腔和孔隙的物体进行强力清洗，以彻底清除喷油嘴顽固积碳。",
    ),
    ("1）Put the externally cleaned injector on the cleaning bracket in the cleaning tank;", "1）将外部已清洁的喷油嘴放在清洗槽内的清洗支架上；"),
    (
        "2）Add an appropriate amount of cleaning agent into the ultrasonic cleaning machine (generally, the level of cleaning agent should be some 20mm above the needle valve of the injector);",
        "2）向超声波清洗机中加入适量清洗剂（一般清洗剂液位应高于喷油嘴针阀约20mm）；",
    ),
    ("3）Connect the pulse signal lines with injectors properly;", "3）正确连接喷油嘴脉冲信号线；"),
    (
        "4）Select [Ultrasonic cleaning] on the main interface and set the time according to the demand (the default time is 600s), as shown in the below figure. Press [Start] button to start the cleaning.",
        "4）在主界面选择【超声波清洗】，按需设置时间（默认600秒），见下图，按【开始】按钮启动清洗。",
    ),
    ("Figure 5.8", "图5.8"),
    ("5）The system will step automatically when the cleaning ends.", "5）清洗结束后系统自动停止。"),
    (
        "6）Take the injectors out of the cleaning tank and wipe off the cleaning agent on them with a soft cloth to prepare for the next operation.",
        "6）将喷油嘴从清洗槽取出，用软布擦去表面清洗剂，准备下一项操作。",
    ),
    (
        "1) It is strictly forbidden to open the ultrasonic system under the circumstances that cleaning agent has not been added into the ultrasonic cleaning pool. Otherwise, the ultrasonic equipment may be damaged easily.",
        "1）严禁在未向超声波清洗槽加入清洗剂的情况下启动超声波系统，否则超声波设备极易损坏。",
    ),
    (
        "2) It is strictly prohibited to immerse the pulse signal line connector along with the injector into the ultrasonic pool for cleaning. Otherwise, the pulse signal line connector can be damaged easily.",
        "2）严禁将脉冲信号线接插件与喷油嘴一起浸入超声波清洗槽清洗，否则接插件极易损坏。",
    ),
    (
        "Leakage test is to detect the leakage of the needle valve of the injector under system pressure and to detect whether the injector is dribbling.",
        "泄漏检测用于在系统压力下检测喷油嘴针阀是否泄漏，以及喷油嘴是否滴漏。",
    ),
    ("Method and Steps (See 5.6 Uniformity/Spray for Installation Method)", "方法与步骤（安装方法见5.6均匀性/雾化检测）"),
    (
        "1）Before leakage test, if there is test solution in the transparent tube, press the fuel drain button on the control panel to drain the solution out of the transparent tube.",
        "1）泄漏检测前，如透明管内有检测液，请按控制面板排油按钮，将检测液从透明管排出。",
    ),
    (
        "2）Select [Leakage Test] on the main interface and press [Start] button. The system will start to work. Judge the leakproofness of the injector by observing whether the injector is dribbling or not. Typically, the dribbling should be no more than one drop in one minute (or technically). The time set in the system is 60 seconds by default, and the pressure value set for leakage test function should be 10% higher than that set by the manufacturer. At this time, the pressure can be adjusted via [Pressurization○+ ] and [Depressurization○- ] buttons, as shown in the below figure.",
        "2）在主界面选择【泄漏检测】并按【开始】按钮，系统开始工作。通过观察喷油嘴是否滴漏判断密封性。通常每分钟滴漏不应超过一滴（或按技术标准）。系统默认时间为60秒，泄漏检测设定压力值应比厂家设定值高10%。此时可通过【加压○+】和【减压○-】按钮调节压力，见下图。",
    ),
    ("Figure 5.9", "图5.9"),
    ("3）The system will stop automatically after completion of the test.", "3）检测完成后系统自动停止。"),
    (
        "Uniformity test is to detect whether the differences between the injection quantities of injectors meet the requirements or within the specified error range under the same working conditions of the injectors on the same vehicle. The test can reflect the electrical characteristics and the change in orifice diameter of the injectors, as well as the combine effects of the blockage and other factors on the injector. Sprayability test is to detect the atomization performance of injectors by observing the injection condition and atomization of injectors when operating under certain working conditions.",
        "均匀性检测用于判断同一车辆各喷油嘴在相同工况下喷油量差异是否符合要求或在规定误差范围内，可反映喷油嘴电气特性、喷孔直径变化，以及堵塞等因素的综合影响。雾化性检测通过观察喷油嘴在特定工况下的喷射状态与雾化情况，检测喷油嘴雾化性能。",
    ),
    ("5.6.1 Installation Method and Test Steps for Injectors", "5.6.1 喷油嘴安装方法及检测步骤"),
    (
        "1) Select the appropriate connector according to the type of injector, install the sealing ring (check if the sealing ring of the connector is in good condition) and then mount the connector with sealing ring at the corresponding coupling element below the fuel distributor.",
        "1）根据喷油嘴类型选择合适接头，安装密封圈（检查接头密封圈是否完好），然后将带密封圈的接头安装到燃油分配器下方对应耦合元件上。",
    ),
    (
        "2) Install the injector in the forward direction (apply a little lubricant on the \"O\" ring of the injector).",
        "2）正向安装喷油嘴（在喷油嘴O型圈上涂抹少量润滑剂）。",
    ),
    (
        "3) Adjust the screws to fix the fuel distributor and injector assembly in the oil filler hole of the upper cover and tighten the compression screws on both sides evenly. The installation diagram is shown in Figure 5.10.",
        "3）调节螺钉，将燃油分配器与喷油嘴总成固定在上盖注油孔中，并均匀拧紧两侧压紧螺钉。安装示意图见图5.10。",
    ),
    ("Figure 5.10", "图5.10"),
    (
        "1-Fuel distributor compression screw; 2-Fuel distributor assembly; 3-Connector; 4-Injector.",
        "1-燃油分配器压紧螺钉；2-燃油分配器总成；3-接头；4-喷油嘴。",
    ),
    (
        "4) Connect the pulse signal line of the injector properly. If there is test solution in the fuel tube, press the fuel drain button on the control panel to drain the test solution out of the transparent tube.",
        "4）正确连接喷油嘴脉冲信号线。如燃油管内有检测液，请按控制面板排油按钮，将检测液从透明管排出。",
    ),
    (
        "5) Click [Uniformity/Spray] button on the main interface as shown in Figure 5.11, set corresponding operating parameters and then press [Start] button to start the test (Note: Press or release the fuel drain button during the operation to drain or stop draining oil); the system pressure can be adjusted by [Pressurization○+ ] and [Depressurization○- ] buttons on the control panel.",
        "5）在主界面点击【均匀性/雾化】按钮（见图5.11），设置相应工况参数后按【开始】按钮启动检测（注意：操作过程中可按放排油按钮进行排油或停止排油）；可通过控制面板【加压○+】和【减压○-】按钮调节系统压力。",
    ),
    (
        "6) The system will stop automatically after the test is completed.",
        "6）检测完成后系统自动停止。",
    ),
    ("Figure 5.11", "图5.11"),
    ("5.6.2 Reverse flushing", "5.6.2 反向冲洗"),
    (
        "Reverse flushing is limited to the top-supply injector by connecting the reverse flushing connector under Uniformity/Spray. Test solution enters from the outlet of the injector and flows out from the inlet during reverse flushing. Reverse flushing can wash away the dirt inside the injector and dirt attached to the filter.",
        "反向冲洗仅适用于顶部供油喷油嘴，在均匀性/雾化检测项下连接反向冲洗接头。反向冲洗时，检测液从喷油嘴出油口进入，从进油口流出，可冲掉喷油嘴内部及滤网附着污物。",
    ),
    (
        "1）Find the reverse flushing connector (and select a supporting \"O\" ring to install it on the connector) installed below the fuel distributor;",
        "1）找到安装在燃油分配器下方的反向冲洗接头（并选择配套O型圈安装到接头上）；",
    ),
    (
        "2）Install the top-supply injector in reverse direction (outlet up, inlet down);",
        "2）反向安装顶部供油喷油嘴（出油口朝上，进油口朝下）；",
    ),
    (
        "3）Select the corresponding coupling element under the injector according to the shape of injector",
        "3）根据喷油嘴形状选择喷油嘴下方对应耦合元件；",
    ),
    (
        "4）Adjust the screws to fix the fuel distributor and injector assembly in the oil filler hole of the upper cover according to the height of injector and tighten the compression screws on both sides evenly, as shown in Figure 5.12;",
        "4）根据喷油嘴高度调节螺钉，将燃油分配器与喷油嘴总成固定在上盖注油孔中，并均匀拧紧两侧压紧螺钉，见图5.12；",
    ),
    (
        "5）It is recommended to press the fuel drain button on the control panel to drain the remaining fuel in the transparent tube to avoid overflow of test solution before reverse flushing.",
        "5）反向冲洗前建议按控制面板排油按钮，排出透明管内残余燃油，避免检测液溢出。",
    ),
    (
        "6）Connect the pulse signal line of the injector properly; set the operating parameters and press [Start] button to execute reverse flushing function; the system pressure can be adjusted by [Pressurization ○+ ] and [Depressurization ○- ] buttons on the control panel during reverse flushing.",
        "6）正确连接喷油嘴脉冲信号线；设置工况参数后按【开始】按钮执行反向冲洗；反向冲洗过程中可通过控制面板【加压○+】和【减压○-】按钮调节系统压力。",
    ),
    (
        "7）The system will stop automatically after completion of the cleaning.",
        "7）清洗完成后系统自动停止。",
    ),
    ("Figure 5.12", "图5.12"),
    (
        "1-Fuel distributor 2-Knurled nut 3-Adjustable screw 4-Couplers 5-Injector 6-O-ring 7-Reverse adaptor 8- Compression screw",
        "1-燃油分配器；2-滚花螺母；3-调节螺钉；4-耦合器；5-喷油嘴；6-O型圈；7-反向冲洗接头；8-压紧螺钉",
    ),
    (
        "Injection quantity is used to detect the amount of fuel injected normally by injectors for 20 seconds and then determine if it is consistent with the injection quantity of standard injectors (or within its error range), referring to the relevant technical manual of the injector). The change or deviation of the change reflects the change (wear) in the orifice diameter or blockage of the injector, eliminating interference due to changes in electrical parameters of the injectors.",
        "喷油量检测用于检测喷油嘴在20秒内正常喷射的燃油量，并判断其是否与标准喷油嘴喷油量一致（或在误差范围内，参考喷油嘴相关技术手册）。喷油量变化或偏差反映喷孔直径变化（磨损）或堵塞情况，可排除喷油嘴电气参数变化带来的干扰。",
    ),
    (
        "Method and Steps (See 5.6 Uniformity/Spray for Installation Method)",
        "方法与步骤（安装方法见5.6均匀性/雾化检测）",
    ),
    (
        "1）Before the test, if there is test solution in the transparent tube, press the fuel drain button on the control panel to drain the solution out of the transparent tube.",
        "1）检测前，如透明管内有检测液，请按控制面板排油按钮，将检测液从透明管排出。",
    ),
    (
        "2）Select [Injection Quantity] on the main interface and press [Start] button. The system will start to work. At this time, the pressure can be adjusted via [Pressurization○+ ] and [Depressurization○- ] buttons, as shown below figure.",
        "2）在主界面选择【喷油量】并按【开始】按钮，系统开始工作。此时可通过【加压○+】和【减压○-】按钮调节压力，见下图。",
    ),
    ("Figure 5.13", "图5.13"),
    ("3）The system will stop automatically after completion of test.", "3）检测完成后系统自动停止。"),
    (
        "Automatic cleaning test includes the above-mentioned several test methods (15-second constant injection fuel injection amount test, idle speed, medium speed, high speed, variable acceleration and deceleration, variable pulse width test). This function can simulate various working conditions of the engine more realistically and comprehensively, and can comprehensively test various performance parameters of the injector.",
        "自动清洗检测包含上述多种检测方法（15秒常喷喷油量检测、怠速、中速、高速、变速加减速、变脉宽检测等）。此功能可更真实、全面地模拟发动机各种工况，综合检测喷油嘴各项性能参数。",
    ),
    (
        "1）Before the test, if there is test solution in the transparent tube, press the fuel drain button to drain the solution out of the transparent tube.",
        "1）检测前，如透明管内有检测液，请按排油按钮将检测液从透明管排出。",
    ),
    (
        "2）Select [Auto Mode] on the main interface and then select a test mode. The default mode is mode 1 (See \"Flow Chart of Auto Mode\" for details on modes). Press [Start] button to start the test.",
        "2）在主界面选择【自动模式】，然后选择检测模式。默认模式为模式1（模式说明见“自动模式流程图”）。按【开始】按钮启动检测。",
    ),
    ("Figure 5.14", "图5.14"),
    (
        "3）During system operation, the pressure can be adjusted via [Pressurization○+ ] and [Depressurization○- ] buttons.",
        "3）系统运行过程中，可通过【加压○+】和【减压○-】按钮调节压力。",
    ),
    (
        "4）When the test is over, the buzzer sounds and the equipment will automatically stop.",
        "4）检测结束时蜂鸣器鸣响，设备自动停止。",
    ),
    ("Flow Chart of Auto Mode", "自动模式流程图"),
    (
        "There are 3 modes for Auto Mode: Mode 1, Mode 2 and Mode 3. Mode 1 and Mode 2 are shown below; while Mode 3 is to run Mode 2 after running Mode 1.",
        "自动模式共有3种：模式1、模式2和模式3。模式1和模式2见下图；模式3为先运行模式1再运行模式2。",
    ),
    ("Mode 1 of Auto Mode", "自动模式—模式1"),
    ("Set the pressure: start the fuel pump to adjust pressure to set value", "设定压力：启动油泵将压力调节至设定值"),
    ("Inject normally for 20s", "正常喷射20秒"),
    ("Observe for 15s, oil drain for 10s", "观察15秒，排油10秒"),
    ("RMP: 650rmp, pulse width: 3ms, counts: 2000", "转速：650rpm，脉宽：3ms，次数：2000"),
    ("Observe idle conditions", "观察怠速工况"),
    ("RMP: 4500rmp, pulse width: 5ms, counts: 1700", "转速：4500rpm，脉宽：5ms，次数：1700"),
    ("Observe medium speed conditions", "观察中速工况"),
    ("RMP: 6000rmp, pulse width: 3ms, counts: 2000", "转速：6000rpm，脉宽：3ms，次数：2000"),
    ("Observe high speed conditions", "观察高速工况"),
    (
        "Observe injection angle and atomization degree",
        "观察喷射角度与雾化程度",
    ),
    (
        "Observe blockage and dripping, normal injection test ended",
        "观察堵塞与滴漏，正常喷射检测结束",
    ),
    ("End", "结束"),
    ("Mode 2 of Auto Mode", "自动模式—模式2"),
    ("Start the device and adjust the pressure to an appropriate value", "启动设备并将压力调节至合适值"),
    ("Speed change test 3 times", "变速检测3次"),
    ("Accelerate from 350rpm to 6000rpm (step: 50rpm)", "从350rpm加速至6000rpm（步进：50rpm）"),
    ("Decelerate from 6000rpm to 350rpm (step: 50rpm)", "从6000rpm减速至350rpm（步进：50rpm）"),
    ("Duty ratio decrease by 1/3 per time at a certain speed", "在某一转速下占空比每次递减1/3"),
    (
        "Speed change test, 3 cycles Accelerate from 350rpm to 6000rpm with duty ratio changing by 1/3 and step by 50rpm Decelerate from 6000rpm to 350rpm with duty ratio changing by 1/3 and step by 50rpm",
        "变速检测，3个循环：从350rpm加速至6000rpm，占空比按1/3变化，步进50rpm；从6000rpm减速至350rpm，占空比按1/3变化，步进50rpm",
    ),
    ("Note:", "注意："),
    ("This function is only used for EFI operation mode.", "此功能仅用于 EFI 工作模式。"),
    (
        "After a period of use, the oil supply system of the engine may be blocked or become clogged due to buildup of dust and impurities in fuel channel. In addition, the carbon deposits and gum made by combustion can easily adhere to the injectors, inlet and outlet ports, inlet and outlet hoses, throttle and combustion chamber. So the fuel supply system, combustion chamber and injectors of the engine must be cleaned in time. On-vehicle clean is a labor-saving and time-saving solution.",
        "使用一段时间后，发动机供油系统可能因燃油通道中灰尘和杂质堆积而堵塞。此外，燃烧产生的积碳和胶质容易附着在喷油嘴、进油口、出油口、进出油管、节气门及燃烧室上。因此必须及时清洗发动机供油系统、燃烧室和喷油嘴。免拆清洗是一种省时省力的解决方案。",
    ),
    ("5.9.1 Procedures", "5.9.1 操作步骤"),
    (
        "1) Please check if the liquid inside the fuel tank is test solution or detergent before performing On-vehicle clean. If test solution is in the tank, it is necessary to replace it with detergent. Drain the test solution inside fuel tank into a pre-prepared container. If the drained test solution contains lots of impurities and cannot be reused, please dispose it in a proper way and then add a small amount of test solution to clean the fuel tank. If the drained test solution is relatively clean, please store it for later use.",
        "1）进行免拆清洗前，请检查油箱内液体是检测液还是清洗剂。如为检测液，须更换为清洗剂。将油箱内检测液排入预先准备的容器。如排出的检测液杂质较多不可再用，请妥善处置，然后加入少量检测液清洗油箱。如排出的检测液较清洁，请妥善保存备用。",
    ),
    (
        "2) Blend the detergent with the fuel at a certain proportion, and fill the mixture into the fuel tank. (Please refer to the cleaning agent's user manual for specific proportions.) Refer to the following table for filling amount:",
        "2）按一定比例将清洗剂与燃油混合后注入油箱。（具体比例请参考清洗剂使用说明。）加注量参考下表：",
    ),
    ("No. of cylinders", "气缸数"),
    ("6 cylinders", "6缸"),
    ("Amount", "用量"),
    ("about 1500ml", "约1500ml"),
    (
        "3) Connect the pipeline of Injector Cleaner & Tester with that of the vehicle. For details, see \"5.9.2 Line Connection\".",
        "3）将喷油嘴清洗检测仪管路与车辆管路连接，详见“5.9.2 管路连接”。",
    ),
    ("Figure 5.15 Installation diagram", "图5.15 安装示意图"),
    ("1- Fuel-return hose from engine; 2- Fuel-inlet hose to engine", "1-发动机回油管；2-发动机进油管"),
    (
        "4) As shown in the following figures: Choose [On-vehicle clean] function on main menu, set the time, press [Start] button and then start the engine for cleaning. Refer to the technical requirements of various vehicle models to adjust the pressure via [Pressurization ○+ ] and [Depressurization ○- ] buttons. Press [Stop] button at any time to stop the cleaning.",
        "4）如下图所示：在主菜单选择【免拆清洗】功能，设置时间后按【开始】按钮，然后启动发动机进行清洗。参照各车型技术要求，通过【加压○+】和【减压○-】按钮调节压力。随时可按【停止】按钮终止清洗。",
    ),
    ("Figure 5.16", "图5.16"),
    ("5.9.2 Line Connection", "5.9.2 管路连接"),
    (
        "There are two cases for the line connection: one is the connection with fuel-return hose and the other is the connection without the fuel-return hose.",
        "管路连接分两种情况：一种带发动机回油管连接，另一种不带发动机回油管连接。",
    ),
    ("Connection with fuel-return hose", "带发动机回油管连接"),
    (
        "1) Disconnect the fuel supply hoses (C, D) and fuel return hoses (A, B) of the engine fuel system (when disconnecting the fuel line connectors, cover them with towels). Choose proper connectors and connect them to the B end and C end separately, and then connect the other ends to corresponding return hose and outlet hose of the device. See Figure 5.17.",
        "1）断开发动机燃油系统的供油管（C、D）和回油管（A、B）（断开燃油管路接头时用毛巾盖住）。选择合适接头分别连接B端和C端，再将另一端分别连接至设备对应回油管和出油管，见图5.17。",
    ),
    (
        "2) Connect the two other disconnected ends (A, D) with a proper hose, or remove the fuel pump fuse, or disconnect the power cable of engine fuel pump.",
        "2）用合适软管连接另外断开的两端（A、D），或拔下油泵保险丝，或断开发动机油泵电源线。",
    ),
    ("to 1 of fig. 5.15", "连接至图5.15之1"),
    ("to 2 of fig. 5.15", "连接至图5.15之2"),
    ("Figure 5.17 Line Connection 1", "图5.17 管路连接1"),
    (
        "1-CNC Injector Cleaner & Tester; 2-Engine; 3- Fuel-return hose from engine; 4- Fuel-inlet hose to engine; 5-Auto filter; 6-Engine fuel supply hose; 7-Engine fuel return hose; 8-Fuel pump; 9-Fuel tank",
        "1-CNC喷油嘴清洗检测仪；2-发动机；3-发动机回油管；4-发动机进油管；5-燃油滤清器；6-发动机供油管；7-发动机回油管；8-油泵；9-油箱",
    ),
    ("Connection without fuel-return hose", "不带发动机回油管连接"),
    (
        "1) Disconnect the fuel supply hoses (E, F) of the engine fuel system (when disconnecting the fuel line connectors, cover them with towels), and then choose a suitable connector, connect it to the E end and reconnect the fuel outlet hose of the device. Hang the fuel return hose. See Figure 5.18.",
        "1）断开发动机燃油系统供油管（E、F）（断开燃油管路接头时用毛巾盖住），选择合适接头连接E端，并重新连接设备出油管。将回油管挂起，见图5.18。",
    ),
    (
        "2) Block the other end of the disconnected end (F) with a stopper (for fuel pump with fuel return function only) or remove the fuse of fuel pump or disconnect the power cable of engine fuel pump.",
        "2）用堵头堵住断开端（F）的另一端（仅适用于带回油功能的油泵），或拔下油泵保险丝，或断开发动机油泵电源线。",
    ),
    ("Figure 5.18 Line Connection 2", "图5.18 管路连接2"),
    (
        "1-CNC Injector Cleaner & Tester; 2-Engine; 3-Stopper; 4- Fuel-inlet hose to engine; 5-Fuel tank; 6- Fuel filter; 7-Engine fuel supply hose; 8-Fuel pump",
        "1-CNC喷油嘴清洗检测仪；2-发动机；3-堵头；4-发动机进油管；5-油箱；6-燃油滤清器；7-发动机供油管；8-油泵",
    ),
    ("5.9.3 Tidy Up After On-vehicle Clean", "5.9.3 免拆清洗后整理"),
    (
        "1) After the On-vehicle clean is completed, turn off the ignition switch of the vehicle. Restitute the hose connection, start the engine and accelerate properly to check if there is any fuel leakage at the connectors or in the hoses.",
        "1）免拆清洗完成后，关闭车辆点火开关。恢复管路连接，启动发动机并适当加速，检查接头及管路是否有燃油泄漏。",
    ),
    (
        "2) Please clean the fuel tank and the lines of the device with test solution at the end of the on-vehicle cleaning, the specific procedures are: drain the detergent left in the fuel tank first and dispose it according to its cleanliness. Add a small amount of test solution into the fuel tank, connect the fuel outlet hose of the device to the oil return port and power on the device. Select \"Leakage test\" item and press [Start] button to run the device for about 2~3 minutes. After the running has stopped, drain the test solution from fuel tank and dispose the drained liquid according to relevant regulations.",
        "2）免拆清洗结束后，请用检测液清洗设备油箱及管路，具体步骤：先排出油箱内剩余清洗剂并按清洁程度处置；向油箱加入少量检测液，将设备出油管连接至回油口并通电；选择“泄漏检测”项目并按【开始】按钮运行约2~3分钟；运行停止后，排出油箱内检测液并按相关规定处置。",
    ),
    (
        "3) Clean up the site and tidy up the washing machine for later use.",
        "3）清理现场，整理清洗设备以备下次使用。",
    ),
    (
        "When cleaning, care must be taken as the detergent is inflammable. Prepare at least one effective fire extinguisher.",
        "清洗时须注意安全，清洗剂为易燃品，请至少准备一只有效灭火器。",
    ),
    (
        "Be sure that all lines are well connected and there is no leakage before performing cleaning.",
        "清洗前务必确认所有管路连接良好、无泄漏。",
    ),
    (
        "It is recommended to lift and carry the device with a manual or motorized forklift.",
        "建议使用手动或机动叉车搬运设备。",
    ),
    ("1. Handling", "1. 搬运"),
    ("A. Mechanical handling and long-distance transportation are strictly prohibited when unpackaged.", "A. 未包装时严禁机械搬运和长途运输。"),
    (
        "B. When unpackaged, use a soft sling for lifting or handle it manually to avoid scratching the body.",
        "B. 拆包后应使用软吊带吊装或人工搬运，避免划伤机身。",
    ),
    (
        "C. The fluid in the fuel tank should be drained before packaging and handling to prevent spillage due to shaking.",
        "C. 包装搬运前应排空油箱内液体，防止晃动泄漏。",
    ),
    (
        "D. The device should be placed on the base and put in the packaging box before handling. Violent vibrations and impacts should be avoided. The machine should be wrapped with materials similar to plastic bags. After placing it into the packaging box, a type of filling materials, such as foam or sponge, should be added between the device and the inner wall of the packaging box to preventing from being scratched by shaking.",
        "D. 搬运前应将设备放在底座上并装入包装箱，避免剧烈振动和冲击。机器应用塑料袋类材料包裹，放入包装箱后，应在设备与箱壁之间填充泡沫或海绵等缓冲材料，防止晃动划伤。",
    ),
    ("E. The maximum angle of inclination for the device must not exceed 45°. Keep upright!", "E. 设备倾斜角度不得超过45°，应保持直立！"),
    ("2. Storing", "2. 存放"),
    ("A. The device should be stored in a dry place that is not exposed to the rain before unpacking.", "A. 开箱前应存放在不淋雨的干燥场所。"),
    (
        "B. The bar machine should be placed in a room that is not exposed to direct sunlight and is well ventilated to avoid rain.",
        "B. 设备应放置在避免阳光直射、通风良好、防雨的房间内。",
    ),
    ("3. Installation Environment", "3. 安装环境"),
    (
        "A. The distance between the equipment and the surrounding walls or other objects should be more than 200mm. The machine should be placed in a ventilated environment with the ambient temperature between 0℃and 45℃. Keep it away from fire.",
        "A. 设备与周围墙壁或其他物体距离应大于200mm。机器应放置在通风良好、环境温度0℃~45℃的环境中，远离火源。",
    ),
    (
        "B. To ensure safe operation, make sure that the power outlet is grounded before switching on the power supply.",
        "B. 为确保安全运行，通电前请确认电源插座已可靠接地。",
    ),
    (
        "If you use the equipment with the power cable being replaced with another one, the level of the power cord should be higher than or equal to that of the original power line.",
        "如更换电源线使用设备，电源线规格等级应不低于原电源线。",
    ),
    ("1. Selection and Replacement of Cleaning Agent and Test Solution", "1. 清洗剂与检测液的选择与更换"),
    (
        "Test solution is used when the device is testing and cleaning agent is used for ultrasonic cleaner. Test solution and cleaning agent are not included in the standard configuration and can be purchased separately.",
        "设备检测时使用检测液，超声波清洗使用清洗剂。检测液和清洗剂均不含在标准配置中，可另行购买。",
    ),
    (
        "When the test solution has been used for a period of time, a lot of impurities will be accumulated in it. Test solution containing lots of dirt cannot be used. Otherwise, the fuel pump and injectors may be blocked. When replacing the test solution, unscrew the plug at the bottom of the machine to drain the solution. It is best to drain the residual liquid and then inject a small amount of clean test solution in order to clean the interior.",
        "检测液使用一段时间后，会积累大量杂质。含杂质过多的检测液不可继续使用，否则可能堵塞油泵和喷油嘴。更换检测液时，拧下机器底部放油塞排液，最好先排尽残液，再注入少量清洁检测液清洗内部。",
    ),
    ("2. Replacement of \"O\" Ring", "2. O型圈更换"),
    (
        "The \"O\" ring will deform after many times of use, which can cause leakage easily. Therefore, it should be replaced frequently.",
        "O型圈多次使用后会变形，容易造成泄漏，应经常更换。",
    ),
    ("3. Replacement of Purification Accessories", "3. 过滤附件更换"),
    (
        "Fuel pump filter is a purification accessory and needs to be replaced regularly. The replacement cycle should be decided according to the service condition and frequency of use. It is recommended to replace the filter once every three months to ensure the normal operation of the system. After replacement, leak tightness test should be carried out to check if there is any leak at the interface.",
        "油泵过滤器为过滤附件，需定期更换。更换周期应根据使用状况和频率确定，建议每三个月更换一次，以确保系统正常运行。更换后应进行密封性检测，检查接口是否泄漏。",
    ),
    (
        "Filter is a purification accessory and needs to be replaced regularly. The replacement cycle should be decided according to the service condition and frequency of use. It is recommended to replace the filter once every three months to ensure the normal operation of the system. After replacement, leak tightness test should be carried out to check if there is any leak at the interface.",
        "过滤器为过滤附件，需定期更换。更换周期应根据使用状况和频率确定，建议每三个月更换一次，以确保系统正常运行。更换后应进行密封性检测，检查接口是否泄漏。",
    ),
    ("1) Replacing filter", "1）更换过滤器"),
    (
        "The equipment filter has been fixed behind the main unit, if you want exchange it you must remove the back board and loose the clamp, take the filter out and exchange a new one.",
        "设备过滤器固定在主机后方，如需更换须拆下后盖板，松开卡箍，取出过滤器并更换新件。",
    ),
    ("2) Replacement of Fuel Pump Filter", "2）更换油泵过滤器"),
    (
        "Fuel pump filter is located in the fuel tank at the bottom of fuel pump. During replacement, it is necessary to remove the fuel pump cover, take off the fuel pump and fuel pump sleeve, unplug the fuel pump filter at the bottom of fuel pump, install a new fuel pump filter, put the fuel pump and fuel pump sleeve back into the fuel tank and put the fuel pump cover back on.",
        "油泵过滤器位于油箱内油泵底部。更换时须拆下油泵盖板，取出油泵及油泵套筒，拔下油泵底部过滤器，安装新过滤器后，将油泵及套筒放回油箱并装回油泵盖板。",
    ),
    ("Figure 5.19", "图5.19"),
    (
        "1 -Filter; 2 -Shunt block; 3- Fuel pump; 4 -Fuel pump bracket; 5- Fuel suction pipe",
        "1-过滤器；2-分流块；3-油泵；4-油泵支架；5-吸油管",
    ),
    ("6.3.1 Precautions", "6.3.1 注意事项"),
    (
        "1）The transparent tubes are made of glass. So, do not place other items around the device to avoid scratching and fragmenting the tubes.",
        "1）透明管为玻璃材质，请勿在设备周围放置其他物品，以免划伤或打碎玻璃管。",
    ),
    (
        "2）Disassembly of the tubes should be performed after the system pressure is displayed as zero.",
        "2）拆卸玻璃管应在系统压力显示为零后进行。",
    ),
    ("3）It must be ensured that the power supply provided is grounded well.", "3）必须确保所提供的电源已良好接地。"),
    (
        "4）Take good care of the machine. If the protective film on the control panel is stained with cleaning agent, please wipe it off in time. Also, keep the pulse signal lines away from the cleaning agent and test solution.",
        "4）请爱护设备。如控制面板保护膜沾有清洗剂，请及时擦除；同时使脉冲信号线远离清洗剂和检测液。",
    ),
    (
        "Blind and imprudent overhauls can lead to the expansion of the fault area of the device, causing difficulties for formal maintenance. When the device is powered on, the electrical system inside the machine contains factors that can cause danger! Careless operations can result in personal injury accidents, which can lead to physical disability and even death in serious accidents.",
        "盲目、不当的检修可能导致故障范围扩大，给正规维修带来困难。设备通电时，机内电气系统存在危险因素！操作不当可能导致人身伤害事故，严重时甚至造成伤残或死亡。",
    ),
    ("6.3.2 Solutions to Common Problems", "6.3.2 常见问题处理"),
    ("1. No response on startup", "1. 开机无反应"),
    (
        "Check if the fuse at the bottom right side of the machine is damaged. If it is damaged, please change it.",
        "检查机器右下方保险丝是否损坏，如损坏请更换。",
    ),
    ("2. Fuel leak at coupling element of fuel distributor", "2. 燃油分配器耦合元件处漏油"),
    (
        "Fuel leaked at coupling element of fuel distributor. Please check whether the \"O\" ring installed matches and if it deformed or damaged. If it does not match or is damaged, please change it. The two adjusting screws should not be too tight, which may also cause the coupling element of fuel distributor to leak fuel.",
        "燃油分配器耦合元件处漏油。请检查所装O型圈是否匹配、是否变形或损坏。如不匹配或已损坏，请更换。两侧调节螺钉不宜过紧，过紧也可能导致耦合元件漏油。",
    ),
    (
        "3. Test solution in the transparent tubes can be drained completely by pressing the fuel drain button on the control panel.",
        "3. 透明管内检测液可通过控制面板排油按钮完全排出。",
    ),
    (
        "Multiple drainages may be required if there is a large amount of test solution in the transparent tube.",
        "如透明管内检测液较多，可能需要多次排油。",
    ),
    ("Manufacturer", "制造商"),
    ("Model", "车型"),
    ("System Pressure (kg/cm2)", "系统压力（kg/cm²）"),
    ("Warranty", "保修条款"),
    (
        "This warranty applies only to users and distributors who have purchased Launch's products through regular procedures.",
        "本保修条款仅适用于通过正规渠道购买元征产品的用户和经销商。",
    ),
    (
        "Launch shall provide a warranty against material or craftsmanship defects for 1 year from the date of delivery on its electronic products. Damages to the device or its components caused by abuses, unauthorized modifications, uses for a purpose other than for which it is intended, or operations not following the manner specified in the manual, etc. are not covered by this warranty.",
        "元征对其电子产品自交付之日起提供一年材料或工艺缺陷保修。因滥用、未经授权改装、用于非预定用途或未按手册规定操作等造成的设备或其部件损坏，不在保修范围内。",
    ),
    ("Disclaimer Statement", "免责声明"),
    ("The above warranty can substitute warranties in any other forms.", "上述保修可替代任何其他形式的保修。"),
    ("Order Notice", "订购须知"),
    (
        "Replaceable and optional parts can be ordered directly from LAUNCH authorized distributors. Your order should include the following information:",
        "可更换件和选配件可直接向元征授权经销商订购。订购时应包含以下信息：",
    ),
    ("Order quantity", "订购数量"),
    ("Part number", "零件号"),
    ("Part name", "零件名称"),
    ("Customer Service Center", "客户服务中心"),
    (
        "For any problem met during the operation, please send email to overseas.service@cnlaunch.com.",
        "操作过程中如有任何问题，请发送邮件至 overseas.service@cnlaunch.com。",
    ),
    (
        "If the device needs to be repaired, please send it back to Launch, and attach the Warranty Card, Product Qualification Certificate, Purchase Invoice and problem description. Launch will maintain and repair the device for free when it is within warranty period. If it is out of warranty, Launch will charge the repair cost and return freight.",
        "如设备需要维修，请寄回元征，并附上保修卡、产品合格证、购货发票及故障说明。保修期内元征将免费维修；保修期外，元征将收取维修费及返程运费。",
    ),
    ("Launch Address:", "元征地址："),
    ("Customer Service Center of LAUNCH TECH CO., LTD.", "深圳市元征科技股份有限公司客户服务中心"),
    (
        "No.4012, Launch Industrial Park, North Wuhe Rd, Bantian Street, Longgang District, Shenzhen, China,",
        "中国深圳市龙岗区坂田街道五和大道北元征工业园4012号，",
    ),
    ("Zip Code: 518129", "邮编：518129"),
    ("Launch Website: https://www.cnlaunch.com", "元征网站：https://www.cnlaunch.com"),
    ("Statement:", "声明："),
    (
        "LAUNCH reserves the rights to make any change to product designs and specifications without notice. The product interface may differ from what is displayed in the manual. Please refer to the actual product for accuracy. We have tried our best to make the descriptions and illustrations in the manual as accurate as possible, and defects are inevitable. If you have any question, please contact local distributor or after-sale service center of LAUNCH. LAUNCH does not bear any responsibility arising from misunderstandings.",
        "元征保留随时更改产品设计和规格的权利，恕不另行通知。产品界面可能与手册显示有所不同，请以实际产品为准。我们已尽力使手册中的描述和图示准确无误，但难免存在疏漏。如有疑问，请联系当地经销商或元征售后服务中心。元征对因误解产生的任何后果不承担责任。",
    ),
    ("CNC605+ User Manual", TITLE_CN),
    ("I. Introduction to Injector Cleaner & Tester", "一、喷油嘴清洗检测仪简介"),
    ("Figure 5.6", "图5.6"),
    ("Figure 5.7", "图5.7"),
    ("ATTENTION！", "注意！"),
    ("blades.", ""),
    ("should be unplugged by hand.", ""),
    ("equipment.", ""),
    ("burns.", ""),
    ("details.", ""),
    ("distributor;", ""),
    ("2.1Structure", "2.1 结构"),
    ("2.2Control Panel", "2.2 控制面板"),
    ("2.2 Control Panel", "2.2 控制面板"),
    ("Control Panel", "控制面板"),
    ("3.1Installation", "3.1 安装"),
    ("3.2Connection", "3.2 连接"),
    ("4.1Preparations", "4.1 准备工作"),
    ("5.5Leakage Test", "5.5 泄漏检测"),
    ("5.8Auto Mode", "5.8 自动模式"),
    ("5.9.1 Procedures", "5.9.1 操作步骤"),
    ("6.3.1 Precautions", "6.3.1 注意事项"),
    ("2. Replacement of \"O\" Ring", "2. O型圈更换"),
    ("No. of cylinders", "气缸数"),
    ("Figure 4.1", "图4.1"),
    ("Figure 4.2", "图4.2"),
    ("Figure 5.1", "图5.1"),
    ("Figure 5.2", "图5.2"),
    ("Figure 5.3", "图5.3"),
    ("Figure 5.4", "图5.4"),
    ("Figure 5.5", "图5.5"),
    ("Figure 5.8", "图5.8"),
    ("Figure 5.9", "图5.9"),
    ("Figure 5.10", "图5.10"),
    ("Figure 5.11", "图5.11"),
    ("Figure 5.12", "图5.12"),
    ("Figure 5.13", "图5.13"),
    ("Figure 5.14", "图5.14"),
    ("Figure 5.15 Installation diagram", "图5.15 安装示意图"),
    ("Figure 5.16", "图5.16"),
    ("Figure 5.17 Line Connection 1", "图5.17 管路连接1"),
    ("Figure 5.18 Line Connection 2", "图5.18 管路连接2"),
    ("Figure 5.19", "图5.19"),
    ("Method and Steps:", "方法与步骤："),
    ("Warning!", "警告！"),
    ("Note:", "注意："),
]


def has_chinese(text: str) -> bool:
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def is_english_fragment(text: str) -> bool:
    stripped = text.strip()
    if not stripped:
        return False
    if has_chinese(stripped):
        return False
    if re.fullmatch(r"[IVXLC]+", stripped):
        return False
    if stripped.isdigit():
        return False
    if stripped in {"LAUNCH", "A"}:
        return False
    ascii_ratio = sum(1 for c in stripped if ord(c) < 128) / len(stripped)
    return ascii_ratio > 0.7 and len(stripped) < 220


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def cleanup_docx(doc: Document) -> None:
    previous = ""
    previous_norm = ""
    for paragraph in list(doc.paragraphs):
        text = paragraph.text.strip()
        if not text:
            continue
        text_norm = normalize(text)
        if text in {'"', '"', '"', '"'}:
            remove_paragraph(paragraph)
            continue
        if text_norm and text_norm == previous_norm:
            remove_paragraph(paragraph)
            continue
        if is_english_fragment(text) and has_chinese(previous):
            remove_paragraph(paragraph)
            continue
        if is_english_fragment(text) and text_norm in previous_norm:
            remove_paragraph(paragraph)
            continue
        previous = text
        previous_norm = text_norm

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in list(cell.paragraphs):
                    text = paragraph.text.strip()
                    if text in {'"', '"'}:
                        remove_paragraph(paragraph)


SORTED_TRANSLATIONS = sorted(TRANSLATIONS, key=lambda item: len(item[0]), reverse=True)
NORM_MAP = {re.sub(r"\s+", "", en).lower(): zh for en, zh in TRANSLATIONS}


def normalize(text: str) -> str:
    return re.sub(r"[^a-z0-9\u4e00-\u9fff]+", "", text.lower())


def translate_text(text: str) -> str:
    if not text.strip():
        return text

    bullet = ""
    core = text
    if core.startswith("\uf06c"):
        bullet = "•\t"
        core = core[1:].lstrip("\t")

    normalized_core = normalize(core)
    if normalized_core in {"i", "ii", "iii"}:
        return text
    if normalized_core in {"attention", "attention！"}:
        return bullet + "注意！"
    if "cnc605usermanual" in normalized_core or normalized_core == normalize(TITLE_EN):
        return bullet + TITLE_CN

    if normalized_core in NORM_MAP:
        return bullet + NORM_MAP[normalized_core]

    for en, zh in SORTED_TRANSLATIONS:
        en_norm = normalize(en)
        if len(en_norm) >= 24 and en_norm in normalized_core:
            return bullet + zh
        if len(normalized_core) >= 24 and normalized_core in en_norm:
            return bullet + zh

    result = core
    for en, zh in SORTED_TRANSLATIONS:
        en_norm = normalize(en)
        if len(en_norm) < 16:
            continue
        pattern = re.escape(en)
        pattern = re.sub(r"\\\s+", r"\\s*", pattern)
        if re.search(pattern, result, flags=re.IGNORECASE):
            result = re.sub(pattern, zh, result, flags=re.IGNORECASE)

    if result != core:
        return bullet + result

    figure_match = re.match(r"Figure\s+([\d.]+)\s*(.*)", core, flags=re.IGNORECASE)
    if figure_match:
        suffix = figure_match.group(2).strip()
        suffix_zh = translate_text(suffix) if suffix else ""
        if suffix_zh and suffix_zh != suffix:
            return bullet + f"图{figure_match.group(1)} {suffix_zh}"
        return bullet + f"图{figure_match.group(1)}"

    section_match = re.match(r"(\d+\.\d+)\s*([A-Za-z][A-Za-z /&-]*)", core)
    if not section_match:
        section_match = re.match(r"(\d+\.\d+)([A-Za-z].*)", core)
    if section_match:
        tail = section_match.group(2).strip()
        tail_zh = NORM_MAP.get(normalize(tail)) or translate_text(tail).lstrip("•\t")
        if tail_zh and tail_zh != tail:
            return bullet + f"{section_match.group(1)} {tail_zh}"

    return text


def translate_docx(src: Path, dst: Path) -> None:
    doc = Document(str(src))
    for p in doc.paragraphs:
        if p.text.strip():
            p.text = translate_text(p.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    cell.text = translate_text(cell.text)
    cleanup_docx(doc)
    doc.save(str(dst))


FONTNAME = "chinahei"
FONT_BUFFER = Path(FONT_PATH).read_bytes()


def fit_text_in_box(page: fitz.Page, rect: fitz.Rect, text: str, fontsize: float) -> None:
    text = text.strip()
    if not text:
        return
    size = fontsize
    expanded = fitz.Rect(rect.x0 - 1, rect.y0 - 1, min(page.rect.width - 20, rect.x1 + 80), rect.y1 + 12)
    while size >= 6:
        rc = page.insert_textbox(
            expanded,
            text,
            fontname=FONTNAME,
            fontsize=size,
            color=(0, 0, 0),
            align=fitz.TEXT_ALIGN_LEFT,
        )
        if rc >= 0:
            return
        size -= 0.4


def should_skip_line(line_text: str) -> bool:
    if not line_text.strip():
        return True
    if line_text in {"LAUNCH", "I", "II", "III"}:
        return True
    if line_text.isdigit() and len(line_text) <= 2:
        return True
    return False


def docx_to_pdf(docx_path: Path, pdf_path: Path) -> None:
    out_dir = pdf_path.parent
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(out_dir),
            str(docx_path),
        ],
        check=True,
        capture_output=True,
        text=True,
    )
    generated = out_dir / f"{docx_path.stem}.pdf"
    if generated != pdf_path and generated.exists():
        generated.replace(pdf_path)


def main() -> None:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    en_docx = OUT_DIR / "CNC605+_User_Manual_EN.docx"
    if not en_docx.exists():
        from pdf2docx import Converter

        cv = Converter(str(SRC_PDF))
        cv.convert(str(en_docx))
        cv.close()
    translate_docx(en_docx, OUT_DOCX)
    docx_to_pdf(OUT_DOCX, OUT_PDF)
    print(f"Created: {OUT_DOCX}")
    print(f"Created: {OUT_PDF}")


if __name__ == "__main__":
    main()
